# coding=utf-8
# @Author: 莫冉
# @Date: 2021-01-20
import docx
from pathlib import Path
import jieba

path = r"F:\实验室\网络小说信息抽取\【AI识别打标资料】作品大纲或脚本等文字资料（1月15）"
file = "《极品战兵》大纲.docx"


# document = docx.Document(Path(path) / file)
#
# for para in document.paragraphs:
#     print(para.text)

# for table in document.tables:
#     for i in range(len(table.rows)):
#         for j in range(len(table.columns)):
#             print(table.cell(i, j).text)


string = "国家暗剑队老大洛羽，因被陷害失去军籍。后来遇上刀疤等人，开始组建自己的组织，收拢天海市各方势力。后来加入龙门派抵抗日本刺客，又进击燕京，成为第一势力。最后前往蒙古拯救父亲并找到宝藏回天海市隐居。"

print(jieba.lcut(string))