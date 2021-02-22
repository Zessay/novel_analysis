# coding=utf-8
# @Author: 莫冉
# @Date: 2021-02-08
import os
import docx
import re
import time
import argparse
import numpy as np
from functools import lru_cache
from typing import Dict, Any, Set
from sklearn.feature_extraction.text import TfidfVectorizer

from novela import logger
import novela.constants as constants
from novela.novel import Comic, Outline, StoryLine, Person
from novela.label import Label, RoleInfo
from novela.text import WordVectorSimilarity, CilinSimilarity, HowNetSimilarity, SentVectorSimilarity
from novela.utils.common import clean_text, load_stopwords
from novela.utils.label_utils import cut_and_remove_stopwords, similarity_between_base_label_and_sents
from novela.utils.label_utils import save_as_excel


def read_document_to_dict(file: str) -> Dict[str, Any]:
    """读取大纲的文档并存储到dict中"""
    # 首先利用正则表达式匹配出小说名
    group = re.findall(r"《(.*?)》", file)
    if len(group) <= 0:
        novel_name = None
    else:
        novel_name = group[0]
    doc_dict = {"小说名": novel_name}     # 用来存储大纲中的内容
    document = docx.Document(file)

    # 由于大纲是按照表格的形式存储的
    # 所以逐个读取
    paras = [p for p in document.paragraphs if p.text.strip()]
    paras_len = len(paras)

    tables = []
    for i, table in enumerate(document.tables):
        if i < paras_len - 1:
            tables.append(table)
        elif i  == paras_len - 1:
            tables.append(document.tables[i:])
            break
        else:
            pass
    # 对于每一个表格以及对应的标题
    for para, table in zip(paras, tables):
        para_title = para.text.strip()   # 表示表格的标题
        # 表示当前属于一个个出场人物的表格
        if isinstance(table, list):
            # 如果table是list
            # 则表明对应出场人物，每一个list对应一个人物
            roles_list = []
            for tb in table:
                role_dict = {}
                for i, row in enumerate(tb.rows):
                    for j, cell in enumerate(row.cells):
                        # 如果是当前行的第一列
                        # 则对应当前行的标题
                        if j == 0:
                            row_title = clean_text(cell.text)
                        else:
                            text = clean_text(cell.text)
                            if text:
                                key, value = text.split("：")
                                if key:
                                    role_dict.update({key:value})
                roles_list.append(role_dict)
            doc_dict.update({para_title:roles_list})
        else:
            table_dict = {}  # 用来存储表格的内容
            row_title = ""
            # 对于表格中的每一行
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    # 如果是当前行的第一列
                    # 则表示当前行的标题
                    if j == 0:
                        # 清洗一下文本
                        row_title = clean_text(cell.text)
                    else:
                        # 如果是`作品资料`这一栏目
                        if "作品资料" in row_title:
                            text = cell.text.strip()
                            if text:
                                text_list = text.split("\n")
                                for t in text_list:
                                    if t:
                                        key, value = t.split("：")
                                        key = clean_text(key)
                                        value_list = value.split()
                                        value = [clean_text(v) for v in value_list]
                                        if len(value) == 1:
                                            value = value[0]
                                        if key:
                                            table_dict.update({key:value})
                        else:
                            # 对于以分行符分段的
                            origin_text = cell.text.strip()
                            if origin_text:
                                text_list = origin_text.split("\n")
                                text_list = [clean_text(t) for t in text_list]
                                if "作品特色" not in row_title:
                                    text = "\n".join(text_list)
                                else:
                                    text = text_list
                                if row_title:
                                    table_dict.update({row_title:text})
            doc_dict[para_title] = table_dict

    return doc_dict


def parse_document(doc_dict: Dict[str, Any], comic: Comic):
    """将原始文档中的信息添加到comic对象中"""
    for table_title, table in doc_dict.items():
        if table_title == "小说名":
            comic.name = table
        elif "大纲" in table_title:
            outline = Outline(title=table.get("名称", ""),
                              novel_author=table.get("原小说作者", ""),
                              comic_author=table.get("漫画作者", ""),
                              novel_type=table.get("作品类型", None))
            for k, v in table.items():
                if k == "读者定位":
                    start, end = v.split("-")
                    outline.target_user_start = float(start)
                    outline.target_user_end = float(end)
                elif "作品特色" in k:
                    outline.features = v
                elif "故事背景" in k:
                    outline.background = v
                elif "主线剧情" in k:
                    outline.summarization = v
                else:
                    pass
            comic.add_outline(outline=outline)
        elif "梗概" in table_title:
            storyline = StoryLine(opening=table.get("开篇", None),
                                  first_step=table.get("起", None),
                                  second_step=table.get("承", None),
                                  third_step=table.get("转（反转）", None),
                                  fourth_step=table.get("合", None))
            comic.add_storyline(storyline=storyline)
        elif "出场人物" in table_title:
            # 添加人物
            for i, role_dict in enumerate(table):
                role = Person(appearance=role_dict.get("外表特征", ""),
                              age=float(role_dict.get("年龄", 0.0)),
                              gender=role_dict.get("性别", ""),
                              desire=role_dict.get("欲求", ""))
                for k, v in role_dict.items():
                    if "姓名" in k:
                        role.name = v
                    elif "兴趣" in k:
                        role.interest = v.split("，")
                    elif "特长" in k:
                        role.speciality = v.split("，")
                    elif "性格" in k:
                        role.character = v.split("，")
                    elif "缺点" in k:
                        role.weakness = v.split("，")
                    else:
                        role.other_info.update({k:v})
                if i == 0:
                    comic.first_role = role
                elif i == 1:
                    comic.second_role = role
                else:
                    comic.add_new_role(role)
        else:
            pass


def classify_base_info(comic: Comic, label: Label):
    """
    对基本信息中的一些标签进行分类
    基本信息中包含：是否原创、连载状况、原著作者、漫画编剧、漫画主笔、男频|女频、是否兼容男频|女频
    """
    # 这里根据comic中的outline分析base_info中的标签
    # -- 原著作者、是否原创 --
    if comic.outline.novel_author:
        label.base_info.original_author = comic.outline.novel_author
        label.base_info.source.value = label.base_info.source.enum_class.NOVEL.display_name
    else:
        label.base_info.source.value = label.base_info.source.enum_class.YC.dislplay_name

    # -- 漫画编剧 --
    if comic.outline.comic_author:
        label.base_info.comic_script_writer = comic.outline.comic_author


@lru_cache(maxsize=1)
def get_story_words_and_sentences(comic: Comic, stopwords: Set[str]):
    """
    获取故事中的有效单词和语句
    :param comic: Comic对象，表示存在的漫画类
    :param stopwords: Set[str]，表示停用词集合
    :return:
    """
    outline_sents_dict = comic.outline.get_sentences()
    outline_sents = []
    for key, value in outline_sents_dict.items():
        if isinstance(value, list):
            outline_sents.extend(value)
        elif isinstance(value, str):
            outline_sents.append(value)
        else:
            pass

    # 获取storyline中的句子
    storyline_sents = comic.storyline.get_sentences()

    sent_strings = outline_sents + storyline_sents
    sentences = [" ".join(cut_and_remove_stopwords(sent, stopwords)) for sent in sent_strings]

    # 定义并训练TFIDF模型
    # 设置token_pattern防止忽略长度为1的单词
    tfidf_model = TfidfVectorizer(token_pattern=r"(?u)\b\w+\b")
    tfidf_model.fit(sentences)

    # 将所有句子拼接成一个长句子
    all_sentence = ""
    for sent in sentences:
        all_sentence += " " + sent
    all_sentence = all_sentence.strip()
    # 将句子中的所有单词转化为TFIDF值
    sparse_value = tfidf_model.transform([all_sentence])   # 结果是一个稀疏矩阵
    vocabulary = tfidf_model.get_feature_names()           # 获取所有单词的词表

    # 获取所有有效的单词以及对应的Tfidf值
    indexes = np.nonzero(sparse_value)[1]
    valid_words, valid_tfidf = [], []
    for ind in indexes:
        valid_words.append(vocabulary[ind])
        valid_tfidf.append(sparse_value[0, ind])

    return sent_strings, valid_words, valid_tfidf, tfidf_model


def classify_story_info(comic: Comic,
                        label: Label,
                        stopwords: Set[str],
                        sim_word2vector: WordVectorSimilarity,
                        sim_hownet: HowNetSimilarity,
                        sim_cilin: CilinSimilarity,
                        sim_sent2vector: SentVectorSimilarity):
    """
    对故事信息中的一些标签进行分类
    """
    # 首先获取故事信息中所有有效的单词和语句
    sent_strings, sent_words, *_ = get_story_words_and_sentences(comic, stopwords)

    # 计算单词和标签之间的相似度
    # --- 主要情节第一层和次要情节第一层 ---
    storyplot_enum_names, storyplot_values = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                                     sent_strings=sent_strings,
                                                                                     base_label=label.story_info.major_storyplot_first,
                                                                                     sim_word2vector=sim_word2vector,
                                                                                     sim_hownet=sim_hownet,
                                                                                     sim_cilin=sim_cilin,
                                                                                     sim_sent2vector=sim_sent2vector,
                                                                                     return_counts=2)
    # 主要情节和次要情节标签
    label.story_info.major_storyplot_first.value = storyplot_values[0]
    label.story_info.minor_storyplot_first.value = storyplot_values[1]

    # --- 主要情节第二层和次要情节第二层 ---
    label.story_info.major_storyplot_first.enum_class = storyplot_enum_names[0]
    label.story_info.minor_storyplot_second.enum_class = storyplot_enum_names[1]
    ## 分别计算两个第二层的值
    _, major_storyplot_second_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                              sent_strings=sent_strings,
                                                                              base_label=label.story_info.major_storyplot_second,
                                                                              sim_word2vector=sim_word2vector,
                                                                              sim_hownet=sim_hownet,
                                                                              sim_cilin=sim_cilin,
                                                                              sim_sent2vector=sim_sent2vector)
    _, minor_storyplot_second_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                              sent_strings=sent_strings,
                                                                              base_label=label.story_info.minor_storyplot_second,
                                                                              sim_word2vector=sim_word2vector,
                                                                              sim_hownet=sim_hownet,
                                                                              sim_cilin=sim_cilin,
                                                                              sim_sent2vector=sim_sent2vector)
    label.story_info.major_storyplot_second.value = major_storyplot_second_value
    label.story_info.minor_storyplot_second.value = minor_storyplot_second_value

    # --- 故事时间 ---
    _, story_time_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                  sent_strings=sent_strings,
                                                                  base_label=label.story_info.story_time,
                                                                  sim_word2vector=sim_word2vector,
                                                                  sim_hownet=sim_hownet,
                                                                  sim_cilin=sim_cilin,
                                                                  sim_sent2vector=sim_sent2vector)
    label.story_info.story_time.value = story_time_value

    # --- 故事文化 ---
    _, story_culture_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                     sent_strings=sent_strings,
                                                                     base_label=label.story_info.story_culture,
                                                                     sim_word2vector=sim_word2vector,
                                                                     sim_hownet=sim_hownet,
                                                                     sim_cilin=sim_cilin,
                                                                     sim_sent2vector=sim_sent2vector)
    label.story_info.story_culture.value = story_culture_value

    # --- 特殊时空 ---
    _, special_space_time_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                          sent_strings=sent_strings,
                                                                          base_label=label.story_info.special_space_time,
                                                                          sim_word2vector=sim_word2vector,
                                                                          sim_hownet=sim_hownet,
                                                                          sim_cilin=sim_cilin,
                                                                          sim_sent2vector=sim_sent2vector)
    label.story_info.special_space_time.value = special_space_time_value

    # --- 故事空间 ---
    _, story_space_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                   sent_strings=sent_strings,
                                                                   base_label=label.story_info.story_space,
                                                                   sim_word2vector=sim_word2vector,
                                                                   sim_hownet=sim_hownet,
                                                                   sim_cilin=sim_cilin,
                                                                   sim_sent2vector=sim_sent2vector)
    label.story_info.story_space.value = story_space_value

    # --- 内容风格 ---
    _, content_style_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                     sent_strings=sent_strings,
                                                                     base_label=label.story_info.content_style,
                                                                     sim_word2vector=sim_word2vector,
                                                                     sim_hownet=sim_hownet,
                                                                     sim_cilin=sim_cilin,
                                                                     sim_sent2vector=sim_sent2vector)
    label.story_info.content_style.value = content_style_value

    # --- 特殊设定 ---
    # List型
    _, special_setting_values = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                        sent_strings=sent_strings,
                                                                        base_label=label.story_info.special_setting,
                                                                        sim_word2vector=sim_word2vector,
                                                                        sim_hownet=sim_hownet,
                                                                        sim_cilin=sim_cilin,
                                                                        sim_sent2vector=sim_sent2vector,
                                                                        return_counts=2)
    label.story_info.special_setting.value = special_setting_values

    # --- 故事套路 ---
    # List型
    _, story_routine_values = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                      sent_strings=sent_strings,
                                                                      base_label=label.story_info.story_routine,
                                                                      sim_word2vector=sim_word2vector,
                                                                      sim_hownet=sim_hownet,
                                                                      sim_cilin=sim_cilin,
                                                                      sim_sent2vector=sim_sent2vector,
                                                                      return_counts=2)
    label.story_info.story_routine.value = story_routine_values


def get_role_words_and_sentences(role: Person, stopwords: Set[str]):
    """获取角色信息中的单词以及语句"""
    sent_words, sent_strings = [], []
    if role.appearance:
        sent_words.extend(cut_and_remove_stopwords(role.appearance, stopwords=stopwords))
        sent_strings.append(role.appearance)
    if role.desire:
        sent_words.extend(cut_and_remove_stopwords(role.desire, stopwords=stopwords))
        sent_strings.append(role.desire)
    if role.interest:
        sent_words.extend(role.interest)
        sent_strings.append("，".join(role.interest))
    if role.character:
        sent_words.extend(role.character)
        sent_strings.append("，".join(role.character))
    if role.speciality:
        sent_words.extend(role.speciality)
        sent_strings.append("，".join(role.speciality))
    if role.weakness:
        sent_words.extend(role.weakness)
        sent_strings.append("，".join(role.weakness))

    return sent_words, sent_strings



def process_role_and_get_role_labels(role: Person,
                                     role_info: RoleInfo,
                                     stopwords: Set[str],
                                     sim_word2vector: WordVectorSimilarity,
                                     sim_hownet: HowNetSimilarity,
                                     sim_cilin: CilinSimilarity,
                                     sim_sent2vector: SentVectorSimilarity):
    """处理单个角色并获取该角色的标签"""
    sent_words, sent_strings = get_role_words_and_sentences(role, stopwords=stopwords)

    # --- 角色物种 ---
    _, role_type_value = similarity_between_base_label_and_sents(sent_words,
                                                                 sent_strings,
                                                                 base_label=role_info.role_type,
                                                                 sim_word2vector=sim_word2vector,
                                                                 sim_hownet=sim_hownet,
                                                                 sim_cilin=sim_cilin,
                                                                 sim_sent2vector=sim_sent2vector)
    role_info.role_type.value = role_type_value

    # --- 角色初始目标 ---
    _, role_target_value = similarity_between_base_label_and_sents(sent_words,
                                                                   sent_strings,
                                                                   base_label=role_info.role_target,
                                                                   sim_word2vector=sim_word2vector,
                                                                   sim_hownet=sim_hownet,
                                                                   sim_cilin=sim_cilin,
                                                                   sim_sent2vector=sim_sent2vector)
    role_info.role_target.value = role_target_value

    # --- 角色职业 ---
    _, role_job_value = similarity_between_base_label_and_sents(sent_words,
                                                                sent_strings,
                                                                base_label=role_info.role_job,
                                                                sim_word2vector=sim_word2vector,
                                                                sim_hownet=sim_hownet,
                                                                sim_cilin=sim_cilin,
                                                                sim_sent2vector=sim_sent2vector)
    role_info.role_job.value = role_job_value

    # --- 角色性格 ---
    _, role_personality_values = similarity_between_base_label_and_sents(sent_words,
                                                                         sent_strings,
                                                                         base_label=role_info.role_personality,
                                                                         sim_word2vector=sim_word2vector,
                                                                         sim_hownet=sim_hownet,
                                                                         sim_cilin=sim_cilin,
                                                                         sim_sent2vector=sim_sent2vector,
                                                                         return_counts=2)
    role_info.role_personality.value = role_personality_values

    # --- 角色外形卖点 ---
    _, role_appearance_value = similarity_between_base_label_and_sents(sent_words,
                                                                       sent_strings,
                                                                       base_label=role_info.role_appearance,
                                                                       sim_word2vector=sim_word2vector,
                                                                       sim_hownet=sim_hownet,
                                                                       sim_cilin=sim_cilin,
                                                                       sim_sent2vector=sim_sent2vector)
    role_info.role_appearance.value = role_appearance_value

    # --- 角色身份卖点 ---
    _, role_identity_value = similarity_between_base_label_and_sents(sent_words,
                                                                     sent_strings,
                                                                     base_label=role_info.role_identity,
                                                                     sim_word2vector=sim_word2vector,
                                                                     sim_hownet=sim_hownet,
                                                                     sim_cilin=sim_cilin,
                                                                     sim_sent2vector=sim_sent2vector)
    role_info.role_identity.value = role_identity_value

    # --- 角色形象卖点 ---
    _, role_figure_value = similarity_between_base_label_and_sents(sent_words,
                                                                   sent_strings,
                                                                   base_label=role_info.role_figure,
                                                                   sim_word2vector=sim_word2vector,
                                                                   sim_hownet=sim_hownet,
                                                                   sim_cilin=sim_cilin,
                                                                   sim_sent2vector=sim_sent2vector)
    role_info.role_figure.value = role_figure_value

    # --- 角色行为卖点 ---
    _, role_behavior_value = similarity_between_base_label_and_sents(sent_words,
                                                                     sent_strings,
                                                                     base_label=role_info.role_behavior,
                                                                     sim_word2vector=sim_word2vector,
                                                                     sim_hownet=sim_hownet,
                                                                     sim_cilin=sim_cilin,
                                                                     sim_sent2vector=sim_sent2vector)
    role_info.role_behavior.value = role_behavior_value

    # --- 角色反差卖点 ---
    _, role_contrast_value = similarity_between_base_label_and_sents(sent_words,
                                                                     sent_strings,
                                                                     base_label=role_info.role_contrast,
                                                                     sim_word2vector=sim_word2vector,
                                                                     sim_hownet=sim_hownet,
                                                                     sim_cilin=sim_cilin,
                                                                     sim_sent2vector=sim_sent2vector)
    role_info.role_contrast.value = role_contrast_value

    return sent_words, sent_strings


def classify_role_info(comic: Comic,
                       label: Label,
                       stopwords: Set[str],
                       sim_word2vector: WordVectorSimilarity,
                       sim_hownet: HowNetSimilarity,
                       sim_cilin: CilinSimilarity,
                       sim_sent2vector: SentVectorSimilarity):
    """对角色信息中的标签进行分类（主要是第一主角和第二主角）"""
    # 首先对第一主角的标签进行分类
    first_words, first_strings = process_role_and_get_role_labels(role=comic.first_role,
                                                                  role_info=label.first_role,
                                                                  stopwords=stopwords,
                                                                  sim_word2vector=sim_word2vector,
                                                                  sim_hownet=sim_hownet,
                                                                  sim_cilin=sim_cilin,
                                                                  sim_sent2vector=sim_sent2vector)

    # 接着对第二主角的标签进行分类
    second_words, second_strings = process_role_and_get_role_labels(role=comic.second_role,
                                                                    role_info=label.second_role,
                                                                    stopwords=stopwords,
                                                                    sim_word2vector=sim_word2vector,
                                                                    sim_hownet=sim_hownet,
                                                                    sim_cilin=sim_cilin,
                                                                    sim_sent2vector=sim_sent2vector)

    # 合并两个角色的单词
    # 得到角色之间关系的标签类别
    sent_words = list(set(first_words + second_words))
    sent_strings = list(set(first_strings + second_strings))

    # --- 主要角色之间的关系（第一层） ---
    relation_first_enum_name, relation_first_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                                             sent_strings=sent_strings,
                                                                                             base_label=label.major_roles_relation_first,
                                                                                             sim_word2vector=sim_word2vector,
                                                                                             sim_hownet=sim_hownet,
                                                                                             sim_cilin=sim_cilin,
                                                                                             sim_sent2vector=sim_sent2vector)
    label.major_roles_relation_first.value = relation_first_value

    # --- 主要角色之间的关系（第二层） ---
    label.major_roles_relation_second.enum_class = relation_first_enum_name
    _, relation_second_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                       sent_strings=sent_strings,
                                                                       base_label=label.major_roles_relation_second,
                                                                       sim_word2vector=sim_word2vector,
                                                                       sim_hownet=sim_hownet,
                                                                       sim_cilin=sim_cilin,
                                                                       sim_sent2vector=sim_sent2vector)
    label.major_roles_relation_second.value = relation_second_value

    # --- 特殊能力 ---
    _, special_ability_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                       sent_strings=sent_strings,
                                                                       base_label=label.special_ability,
                                                                       sim_word2vector=sim_word2vector,
                                                                       sim_hownet=sim_hownet,
                                                                       sim_cilin=sim_cilin,
                                                                       sim_sent2vector=sim_sent2vector)
    label.special_ability.value = special_ability_value

    # --- 外挂 ---
    _, plugin_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                              sent_strings=sent_strings,
                                                              base_label=label.plugin,
                                                              sim_word2vector=sim_word2vector,
                                                              sim_hownet=sim_hownet,
                                                              sim_cilin=sim_cilin,
                                                              sim_sent2vector=sim_sent2vector)
    label.plugin.value = plugin_value


def classify_other_info(comic: Comic,
                        label: Label,
                        stopwords: Set[str],
                        sim_word2vector: WordVectorSimilarity,
                        sim_hownet: HowNetSimilarity,
                        sim_cilin: CilinSimilarity,
                        sim_sent2vector: SentVectorSimilarity):
    """对其他信息中的标签进行分类"""
    # 首先获取文章中的所有有效单词和语句
    sent_strings, sent_words, *_ = get_story_words_and_sentences(comic, stopwords)

    # --- 热门话题 ---
    _, hot_topic_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                 sent_strings=sent_strings,
                                                                 base_label=label.others.hot_topic,
                                                                 sim_word2vector=sim_word2vector,
                                                                 sim_hownet=sim_hownet,
                                                                 sim_cilin=sim_cilin,
                                                                 sim_sent2vector=sim_sent2vector)
    label.others.hot_topic.value = hot_topic_value

    # --- 其他卖点 ---
    _, other_points_value = similarity_between_base_label_and_sents(sent_words=sent_words,
                                                                    sent_strings=sent_strings,
                                                                    base_label=label.others.other_points,
                                                                    sim_word2vector=sim_word2vector,
                                                                    sim_hownet=sim_hownet,
                                                                    sim_cilin=sim_cilin,
                                                                    sim_sent2vector=sim_sent2vector)
    label.others.other_points.value = other_points_value




if __name__ == '__main__':
    # doc_dict = read_document_to_dict(os.path.join(data_path, file))
    # print(doc_dict)
    # # read_test(os.path.join(data_path, file))
    parser = argparse.ArgumentParser()
    parser.add_argument("--source_dir", default="/home/data/corpus", type=str, required=True,
                        help="The directory name of the base path.")
    parser.add_argument("--file_name", default="《极品战兵》大纲.docx", type=str, required=True,
                        help="The file name of the outline.")
    parser.add_argument("--w2v_file", default="/home/models/wordvector/sgns.literature.word.bz2", type=str,
                        help="The path of word2vector file.")
    parser.add_argument("--to_dir", default="/home/results/novel", type=str, required=True,
                        help="The target directory of the result file.")
    parser.add_argument("--to_file", default="novel_label.xlsx", type=str, required=True,
                        help="The target file name.")

    args = parser.parse_args()

    # 首先判断文件夹和文件是否存在
    if not os.path.isdir(args.source_dir):
        raise RuntimeError(f"The parameter `source_dir`: {args.source_dir} is not a valid path.")

    source_file = os.path.join(args.source_dir, args.file_name)
    if not os.path.isfile(source_file):
        raise ValueError(f"There is no file named `source_file`: {source_file}.")

    # 判断目标文件目录是否存在
    if not os.path.isdir(args.to_dir):
        os.makedirs(args.to_dir)
    target_file = os.path.join(args.to_dir, args.to_file)


    # 将文件中的数据转化为Dict型
    logger.info(f"读取小说大纲文件{source_file}，并转化为Comic对象")
    document_dict = read_document_to_dict(file=source_file)
    novel_name = document_dict["小说名"]                     # 小说的名称
    # 转化为Comic对象
    comic = Comic()
    parse_document(document_dict, comic=comic)

    # 加载停用词
    stopwords = load_stopwords(file=constants.STOPWORDS_FILE)

    # 创建相似度计算的对象
    # --- 单词相似度部分 ---
    logger.info("构造单词语义相似度计算对象")
    sim_word2vector = WordVectorSimilarity(w2v_file=args.w2v_file)
    logger.info("构造词林相似度对象")
    sim_cilin = CilinSimilarity(cilin_file=constants.CILIN_FILE)
    logger.info("构造HowNet相似度计算对象")
    sim_hownet = HowNetSimilarity(glossary_file=constants.GLOSSARY_FILE,
                                  sememe_file=constants.WHOLE_DAT)
    # --- 语句相似度部分 ---
    logger.info("构造语句相似度计算对象")
    sim_sent2vector = SentVectorSimilarity(stopwords=stopwords,
                                           word2vec=sim_word2vector.word2vec)

    # 创建空的Label对象
    label = Label()

    start_time = time.time()
    classify_base_info(comic=comic, label=label)
    logger.info(f"得到基本信息的标签，共计用时 {time.time()-start_time} s.")

    start_time = time.time()
    classify_story_info(comic=comic, label=label,
                        stopwords=stopwords,
                        sim_word2vector=sim_word2vector,
                        sim_cilin=sim_cilin,
                        sim_hownet=sim_hownet,
                        sim_sent2vector=sim_sent2vector)
    logger.info(f"得到故事信息的标签，共计用时 {time.time()-start_time} s.")

    start_time = time.time()
    classify_role_info(comic=comic, label=label,
                       stopwords=stopwords,
                       sim_word2vector=sim_word2vector,
                       sim_cilin=sim_cilin,
                       sim_hownet=sim_hownet,
                       sim_sent2vector=sim_sent2vector)
    logger.info(f"得到角色信息的标签，共计用时 {time.time()-start_time} s.")

    start_time = time.time()
    classify_other_info(comic=comic, label=label,
                        stopwords=stopwords,
                        sim_word2vector=sim_word2vector,
                        sim_cilin=sim_cilin,
                        sim_hownet=sim_hownet,
                        sim_sent2vector=sim_sent2vector)
    logger.info(f"得到其他信息的标签，共计用时 {time.time()-start_time} s.")


    logger.info(f"保存文件到{target_file}")
    save_as_excel(to_file=target_file,
                  novel_name=novel_name,
                  label=label)
    logger.info("保存完成！")

